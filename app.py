import os
import io
import base64
import re
import json
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from PIL import Image
import pytesseract
from docx import Document
from openpyxl import Workbook

# ======================================================================
# TESSERACT CONFIGURATION
# IMPORTANT: This path is for local Windows development only.
# It MUST be removed or commented out for deployment on platforms like Vercel.
# ======================================================================
TESSERACT_PATH = r'C:\Users\myatphonesan\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
try:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
except Exception:
    # Allows the app to run locally even if Tesseract isn't immediately configured
    pass

# --- FLASK APP DEFINITION ---
app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET', 'devsecret')
# ----------------------------


def run_ocr_from_bytes(file_bytes: bytes):
    """Run OCR on raw image bytes and return (text, error_message)."""
    try:
        image = Image.open(io.BytesIO(file_bytes))
        
       
        text = pytesseract.image_to_string(image, lang='eng+mya') 
        
        return text, None
    except pytesseract.TesseractNotFoundError:
        return None, "Tesseract executable not found. Please check the TESSERACT_PATH."
    except Exception as e:
        # Include specific Tesseract language errors here if possible, but a general
        # error catch is fine for this example.
        return None, f"Unexpected error during OCR: {e}"


def parse_trade_document(ocr_text: str):
    """
    Performs generic key-value extraction by iterating through lines 
    and splitting on common delimiters like colons, treating the first
    part as the 'key' (row name) and the rest as the 'value'.
    """
    parsed_data = {}
    
    # Split the entire text into lines and clean up blank lines
    lines = [line.strip() for line in ocr_text.split('\n') if line.strip()]

    for i, line in enumerate(lines):
        # 1. Attempt to find a Key-Value separator (e.g., 'Name: John Doe', 'Amount 123.45')
        
        # Pattern 1: Find a colon or equals sign as a clear separator
        # This regex should work with Myanmar characters as well.
        match = re.search(r'(.+?)[:=]\s*(.+)', line)
        if match:
            key = match.group(1).strip()
            value = match.group(2).strip()
            
            # Simple check to ensure the key isn't just a number (this is a simple heuristic)
            if key and not key.isdigit():
                # Clean up the key name (This cleaning might be too aggressive for Myanmar keys, 
                # but we'll keep it simple for now)
                cleaned_key = re.sub(r'[^a-zA-Z0-9\s\u1000-\u109f]', '', key).strip() 
                # \u1000-\u109f covers the basic Myanmar script block
                
                # Use the cleaned key directly (or add an index if duplicates are expected)
                parsed_data[cleaned_key] = value
                continue

        # Pattern 2: Line without a clear separator
        # This captures data that doesn't fit the key:value pattern.
        if line.strip():
             # Using a line number as a fallback key
             key = f"Line {i+1}"
             parsed_data[key] = line.strip()

    return parsed_data


@app.route('/', methods=['GET', 'POST'])
def index():
    extracted_text = None
    structured_data = {}
    structured_data_json = None
    filename = None
    image_data = None
    
    if request.method == 'POST':
        uploaded = request.files.get('file')
        if not uploaded or uploaded.filename == '':
            flash('No file selected', 'warning')
            return redirect(url_for('index'))

        ext = uploaded.filename.rsplit('.', 1)[-1].lower()
        if ext not in ('png', 'jpg', 'jpeg'):
            flash('Unsupported file type. Please upload PNG/JPG/JPEG.', 'danger')
            return redirect(url_for('index'))

        file_bytes = uploaded.read()
        extracted_text, error = run_ocr_from_bytes(file_bytes)
        filename = uploaded.filename
        
        if error:
            flash(error, 'danger')
            return render_template('index.html', tesseract_path=TESSERACT_PATH)

        # --- TDA STEP: PARSE THE DATA ---
        structured_data = parse_trade_document(extracted_text)
        structured_data_json = json.dumps(structured_data) # Convert dict to JSON string for passing to HTML
        
        # Build a data URI to preview the image in the browser
        image_b64 = base64.b64encode(file_bytes).decode('utf-8')
        image_data = f"data:{uploaded.mimetype};base64,{image_b64}"

    return render_template(
        'index.html',
        extracted_text=extracted_text,
        structured_data=structured_data,
        structured_data_json=structured_data_json, # Used to pass data back to /download
        filename=filename,
        image_data=image_data,
        tesseract_path=TESSERACT_PATH,
    )


@app.route('/download', methods=['POST'])
def download():
    text = request.form.get('extracted_text', '')
    structured_data_json = request.form.get('structured_data_json', '{}')
    
    original_file_name_with_ext = request.form.get('filename', 'ocr_result.txt')
    file_type = request.form.get('file_type', 'txt')
    
    if not text:
        flash('Nothing to download', 'warning')
        return redirect(url_for('index'))

    base_name = os.path.splitext(original_file_name_with_ext)[0]
    buf = io.BytesIO()

    if file_type == 'txt':
        # --- TXT FILE GENERATION (Raw Text) ---
        # Include structured data at the top for clarity
        output_text = "--- Structured Data (Key: Value) ---\n"
        try:
            data = json.loads(structured_data_json)
            for key, value in data.items():
                 output_text += f"{key}: {value}\n"
        except json.JSONDecodeError:
            output_text += "Could not parse structured data.\n"
            
        output_text += "\n--- Full Raw OCR Text ---\n"
        output_text += text
        
        buf.write(output_text.encode('utf-8'))
        download_name = f"{base_name}.txt"
        mimetype = 'text/plain'
    
    elif file_type == 'docx':
        # --- WORD (.docx) FILE GENERATION (Raw Text) ---
        document = Document()
        document.add_paragraph("--- Structured Data (Row-by-Row Extraction) ---")
        try:
            data = json.loads(structured_data_json)
            if data:
                table = document.add_table(rows=len(data) + 1, cols=2)
                table.style = 'Table Grid'
                # Add headers
                table.rows[0].cells[0].text = 'Row Name (Key)'
                table.rows[0].cells[1].text = 'Value'
                
                row_index = 1
                for key, value in data.items():
                    table.rows[row_index].cells[0].text = key
                    table.rows[row_index].cells[1].text = str(value)
                    row_index += 1
            else:
                 document.add_paragraph("No structured data extracted.")

        except json.JSONDecodeError:
            document.add_paragraph("Could not parse structured data.")

        document.add_paragraph("\n--- Full OCR Text ---")
        document.add_paragraph(text)
        document.save(buf)
        download_name = f"{base_name}.docx"
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    elif file_type == 'xlsx':
        # --- EXCEL (.xlsx) FILE GENERATION (Structured Data) ---
        wb = Workbook()
        ws = wb.active
        ws.title = "Extracted OCR Data"
        
        try:
            data = json.loads(structured_data_json)
            
            if data:
                # Write Header Row
                ws.append(['Row Name (Key)', 'Value'])
                
                # Write Data Rows
                for key, value in data.items():
                     ws.append([key, value])
            else:
                ws['A1'] = "No structured data extracted."

            # Optional: Add the raw text to a second sheet
            ws_raw = wb.create_sheet(title="Raw OCR Text")
            ws_raw['A1'] = text
            
        except json.JSONDecodeError:
            ws['A1'] = "Error: Could not parse structured data."
            ws['A2'] = "Raw OCR Text:"
            ws['A3'] = text
        
        wb.save(buf)
        download_name = f"{base_name}.xlsx"
        mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

    else:
        flash('Invalid download file type requested.', 'danger')
        return redirect(url_for('index'))

    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name=download_name, mimetype=mimetype)


if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=8501)